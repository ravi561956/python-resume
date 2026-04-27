from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from io import BytesIO
from bs4 import BeautifulSoup


# =============================
# HTML → TEXT
# =============================
def html_to_text(html):
    if not html:
        return ""
    return BeautifulSoup(html, "html.parser").get_text("\n")


# =============================
# HELPER: REMOVE SPACING
# =============================
def set_paragraph_spacing(paragraph, before=5, after=5, line=1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def generate_resume_docx(resume):
    doc = Document()

    # GLOBAL STYLE
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # PAGE MARGINS
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # =============================
    # HEADER
    # =============================
    table = doc.add_table(rows=1, cols=2)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    # LEFT
    p = left.paragraphs[0]
    run = p.add_run(resume.name)
    run.bold = True
    run.font.size = Pt(20)

    p = left.add_paragraph(resume.position)

    # RIGHT
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    current_date = datetime.now().strftime("%d %b %Y")
    p.add_run(f"Date: {current_date}")

    # =============================
    # CONTACT
    # =============================
    doc.add_paragraph(f"Phone: {resume.phone}")
    doc.add_paragraph(f"Email: {resume.email}")
    doc.add_paragraph(f"Address: {resume.address}")

    if resume.website:
        doc.add_paragraph(f"Website: {resume.website}")

    # =============================
    # SUMMARY
    # =============================
    summary = html_to_text(resume.summary)

    if summary:
        doc.add_paragraph("Professional Summary").runs[0].bold = True
        doc.add_paragraph(summary)

    # =============================
    # SKILLS
    # =============================
    skills = []
    for category in resume.skill_categories.all():
        for skill in category.skills.all():
            skills.append(skill.name)

    if skills:
        doc.add_paragraph("Skills").runs[0].bold = True

        table = doc.add_table(rows=1, cols=3)

        col_data = [
            skills[0::3],
            skills[1::3],
            skills[2::3],
        ]

        for i, col in enumerate(col_data):
            cell = table.rows[0].cells[i]
            for skill in col:
                cell.add_paragraph(f"• {skill}")

    # =============================
    # EXPERIENCE
    # =============================
    if resume.journey.all():
        doc.add_paragraph("Work Experience").runs[0].bold = True

        for journey in resume.journey.all():
            for item in journey.excellences.all():
                doc.add_paragraph(item.title).runs[0].bold = True
                doc.add_paragraph(f"{item.company} | {item.date_range}")

                desc = html_to_text(item.description)
                doc.add_paragraph(desc)

    # =============================
    # CERTIFICATIONS
    # =============================
    certs = []
    for profession in resume.professions.all():
        for cert in profession.certifications.all():
            if cert.is_active:
                certs.append(cert.title)

    if certs:
        doc.add_paragraph("Certifications").runs[0].bold = True

        table = doc.add_table(rows=1, cols=2)

        col_data = [
            certs[0::2],
            certs[1::2],
        ]

        for i, col in enumerate(col_data):
            cell = table.rows[0].cells[i]
            for cert in col:
                cell.add_paragraph(f"• {cert}")

    # =============================
    # EDUCATION
    # =============================
    if resume.journeys.all():
        doc.add_paragraph("Education").runs[0].bold = True

        for item in resume.journeys.all():
            if item.is_active:
                desc = html_to_text(item.description)
                doc.add_paragraph(f"{desc} ({item.year})")

    # =============================
    # SAVE
    # =============================
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer